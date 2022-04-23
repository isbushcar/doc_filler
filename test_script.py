import os
import uuid
from typing import Union

import pandas as pd
from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate
from docx.enum.text import WD_BREAK


# variables_table = pd.read_excel('test_table.xlsx')
# print(variables_table)
# print(list(enumerate(variables_table.index)))


def fill_document(
        doc_path: Union[str, os.PathLike],
        variables_table_path: Union[str, os.PathLike],
) -> str:
    variables_table = pd.read_excel(variables_table_path)
    # variables_table.columns = [str(x).lower() for x in variables_table.columns]
    generated_docs = []

    def generate_numbered_names(variables_table_index: pd.Index):
        temp_name = str(uuid.uuid4())
        for index in range(len(variables_table_index)):
            yield f'{temp_name}_{index}.docx'

    file_name = generate_numbered_names(variables_table.index)

    for number, row_index in enumerate(variables_table.index):
        document_to_fill = DocxTemplate(doc_path)
        current_context = {}
        for column in variables_table.columns:
            current_context[column] = variables_table.loc[row_index, column]
        document_to_fill.render(current_context)

        new_name = next(file_name)
        document_to_fill.save(new_name)

        if not number == len(variables_table.index) - 1:
            document_to_fill = Document(new_name)
            document_to_fill.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
            document_to_fill.save(new_name)

        generated_docs.append(new_name)

    for num, doc in enumerate(generated_docs):
        if num == 0:
            base_doc = Composer(Document(doc))
        else:
            base_doc.append(Document(doc))

    combined_doc_name = f"{str(uuid.uuid4())}.docx"
    base_doc.save(combined_doc_name)

    for num, doc in enumerate(generated_docs):
        os.remove(doc)

    return combined_doc_name


fill_document('test.docx', 'test_table.xlsx')
